"""Document Management Services - OCR and AI Summary."""

from __future__ import annotations

import base64
import logging
import os
import subprocess
from typing import Optional

from django.conf import settings

logger = logging.getLogger(__name__)


class OCRService:
    """Service for extracting text from documents using OCR."""
    
    # Supported file types for OCR
    SUPPORTED_IMAGE_TYPES = {'image/png', 'image/jpeg', 'image/jpg', 'image/tiff', 'image/bmp', 'image/gif'}
    SUPPORTED_PDF_TYPES = {'application/pdf'}
    SUPPORTED_DOCX_TYPES = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'  # Also support old .doc format if converted
    }
    
    @classmethod
    def is_supported(cls, mime_type: str) -> bool:
        """Check if the file type is supported for OCR."""
        return mime_type.lower() in cls.SUPPORTED_IMAGE_TYPES | cls.SUPPORTED_PDF_TYPES | cls.SUPPORTED_DOCX_TYPES
    
    @classmethod
    def extract_text(cls, file_path: str, mime_type: str) -> Optional[str]:
        """
        Extract text from a document using OCR.
        
        Args:
            file_path: Path to the file on disk
            mime_type: MIME type of the file
            
        Returns:
            Extracted text or None if extraction failed
        """
        if not os.path.exists(file_path):
            logger.warning(f"OCR: File not found: {file_path}")
            return None
        
        mime_type = mime_type.lower()
        
        try:
            if mime_type in cls.SUPPORTED_PDF_TYPES:
                return cls._extract_from_pdf(file_path)
            elif mime_type in cls.SUPPORTED_IMAGE_TYPES:
                return cls._extract_from_image(file_path)
            elif mime_type in cls.SUPPORTED_DOCX_TYPES:
                return cls._extract_from_docx(file_path)
            else:
                logger.info(f"OCR: Unsupported file type: {mime_type}")
                return None
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return None
    
    @classmethod
    def _extract_from_image(cls, file_path: str) -> Optional[str]:
        """Extract text from an image using pytesseract."""
        try:
            import pytesseract
            from PIL import Image
            
            # Open and process image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary (for better OCR results)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Run OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            # Clean up text
            text = text.strip()
            
            if text:
                logger.info(f"OCR: Extracted {len(text)} characters from image")
                return text
            else:
                logger.info("OCR: No text found in image")
                return None
                
        except ImportError:
            logger.warning("OCR: pytesseract or PIL not installed")
            return cls._extract_with_tesseract_cli(file_path)
        except Exception as e:
            logger.error(f"OCR image extraction failed: {e}")
            return None
    
    @classmethod
    def _extract_from_pdf(cls, file_path: str) -> Optional[str]:
        """Extract text from a PDF using pdf2image and pytesseract."""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)
            
            if not images:
                logger.warning("OCR: No pages found in PDF")
                return None
            
            # Extract text from each page
            all_text = []
            for i, image in enumerate(images):
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Run OCR on page
                page_text = pytesseract.image_to_string(image, lang='eng')
                if page_text.strip():
                    all_text.append(f"--- Page {i + 1} ---\n{page_text.strip()}")
            
            if all_text:
                combined_text = "\n\n".join(all_text)
                logger.info(f"OCR: Extracted {len(combined_text)} characters from PDF ({len(images)} pages)")
                return combined_text
            else:
                logger.info("OCR: No text found in PDF")
                return None
                
        except ImportError:
            logger.warning("OCR: pdf2image or pytesseract not installed")
            return cls._extract_pdf_with_pdftotext(file_path)
        except Exception as e:
            logger.error(f"OCR PDF extraction failed: {e}")
            return None
    
    @classmethod
    def _extract_with_tesseract_cli(cls, file_path: str) -> Optional[str]:
        """Fallback: Extract text using tesseract CLI directly."""
        try:
            result = subprocess.run(
                ['tesseract', file_path, 'stdout'],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
            return None
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning(f"Tesseract CLI fallback failed: {e}")
            return None
    
    @classmethod
    def _extract_pdf_with_pdftotext(cls, file_path: str) -> Optional[str]:
        """Fallback: Extract text from PDF using pdftotext CLI."""
        try:
            result = subprocess.run(
                ['pdftotext', file_path, '-'],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
            return None
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning(f"pdftotext fallback failed: {e}")
            return None
    
    @classmethod
    def _extract_from_docx(cls, file_path: str) -> Optional[str]:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document as DocxDocument
            
            # Open the DOCX file
            doc = DocxDocument(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n--- Tables ---\n\n" + "\n".join(table_texts)
            
            if full_text.strip():
                logger.info(f"OCR: Extracted {len(full_text)} characters from DOCX")
                return full_text.strip()
            else:
                logger.info("OCR: No text found in DOCX")
                return None
                
        except ImportError:
            logger.warning("OCR: python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"OCR DOCX extraction failed: {e}")
            return None


class DocumentSummaryService:
    """Service for generating AI summaries of documents."""
    
    # Maximum characters to process for summary
    MAX_CONTENT_LENGTH = 50000
    
    # Summary prompt template
    SUMMARY_PROMPT = """Please provide a concise summary of the following document.
The summary should:
- Be 2-4 paragraphs long
- Highlight the main points and key takeaways
- Be written in a professional tone
- Include any important dates, names, or figures mentioned

Document Title: {title}

Document Content:
{content}

Summary:"""

    @classmethod
    def generate_summary(cls, content: str, title: str = "Untitled Document") -> str:
        """
        Generate a summary of document content.
        
        This method provides a basic extractive summary. For production use,
        you can integrate with OpenAI, Anthropic, or other LLM APIs.
        
        Args:
            content: The document text content
            title: Document title for context
            
        Returns:
            Generated summary text
        """
        if not content or not content.strip():
            return "No content available for summarization."
        
        # Truncate if too long
        if len(content) > cls.MAX_CONTENT_LENGTH:
            content = content[:cls.MAX_CONTENT_LENGTH] + "..."
        
        # Try LLM-based summary first
        llm_summary = cls._generate_with_llm(content, title)
        if llm_summary:
            return llm_summary
        
        # Fall back to extractive summary
        return cls._generate_extractive_summary(content, title)
    
    @classmethod
    def _generate_with_llm(cls, content: str, title: str) -> Optional[str]:
        """
        Generate summary using an LLM API.
        
        Configure your LLM API credentials in settings:
        - OPENAI_API_KEY for OpenAI
        - ANTHROPIC_API_KEY for Anthropic Claude
        """
        # Try OpenAI first
        openai_key = getattr(settings, 'OPENAI_API_KEY', None)
        if openai_key:
            try:
                import openai
                openai.api_key = openai_key
                
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {
                            "role": "system",
                            "content": "You are a helpful assistant that summarizes documents concisely and professionally."
                        },
                        {
                            "role": "user",
                            "content": cls.SUMMARY_PROMPT.format(title=title, content=content)
                        }
                    ],
                    max_tokens=500,
                    temperature=0.3
                )
                
                summary = response.choices[0].message.content.strip()
                logger.info(f"Generated LLM summary ({len(summary)} chars)")
                return summary
                
            except ImportError:
                logger.info("OpenAI package not installed")
            except Exception as e:
                logger.warning(f"OpenAI summary generation failed: {e}")
        
        # Try Anthropic Claude
        anthropic_key = getattr(settings, 'ANTHROPIC_API_KEY', None)
        if anthropic_key:
            try:
                import anthropic
                client = anthropic.Anthropic(api_key=anthropic_key)
                
                message = client.messages.create(
                    model="claude-3-haiku-20240307",
                    max_tokens=500,
                    messages=[
                        {
                            "role": "user",
                            "content": cls.SUMMARY_PROMPT.format(title=title, content=content)
                        }
                    ]
                )
                
                summary = message.content[0].text.strip()
                logger.info(f"Generated Claude summary ({len(summary)} chars)")
                return summary
                
            except ImportError:
                logger.info("Anthropic package not installed")
            except Exception as e:
                logger.warning(f"Anthropic summary generation failed: {e}")
        
        return None
    
    @classmethod
    def _generate_extractive_summary(cls, content: str, title: str) -> str:
        """
        Generate a simple extractive summary by selecting key sentences.
        This is a fallback when no LLM API is available.
        """
        import re
        
        # Split into sentences
        sentences = re.split(r'[.!?]+', content)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]
        
        if not sentences:
            return "Unable to generate summary - document content too short or unstructured."
        
        # Score sentences (simple heuristics)
        scored_sentences = []
        keywords = set(title.lower().split()) if title else set()
        
        for i, sentence in enumerate(sentences):
            score = 0
            words = sentence.lower().split()
            
            # Position bonus (first and last sentences often important)
            if i < 3:
                score += 2
            if i >= len(sentences) - 2:
                score += 1
            
            # Keyword matching
            keyword_matches = sum(1 for word in words if word in keywords)
            score += keyword_matches * 2
            
            # Length penalty for very short or very long sentences
            if 50 <= len(sentence) <= 200:
                score += 1
            
            # Contains important words
            important_words = {'important', 'key', 'main', 'significant', 'critical', 
                            'therefore', 'conclusion', 'summary', 'result', 'finding'}
            if any(word in words for word in important_words):
                score += 2
            
            scored_sentences.append((score, sentence))
        
        # Sort by score and select top sentences
        scored_sentences.sort(key=lambda x: x[0], reverse=True)
        
        # Take top 5 sentences, but maintain original order
        top_sentences = scored_sentences[:5]
        selected_indices = sorted(
            [sentences.index(s) for _, s in top_sentences if s in sentences]
        )
        
        summary_sentences = [sentences[i] for i in selected_indices]
        
        if not summary_sentences:
            summary_sentences = sentences[:3]  # Fallback to first 3 sentences
        
        summary = ". ".join(summary_sentences)
        if not summary.endswith('.'):
            summary += "."
        
        # Add title context
        if title and title != "Untitled Document":
            summary = f"Summary of '{title}':\n\n{summary}"
        
        return summary


class FileUploadService:
    """Service for processing file uploads (base64 decode, save to disk, OCR)."""

    SUPPORTED_OCR_TYPES = {
        'image/png', 'image/jpeg', 'image/jpg', 'image/tiff', 'application/pdf',
    }

    @classmethod
    def process_data_url(cls, file_url: str, file_name: str, file_type: str,
                         document_identifier: str | None = None,
                         version: object | None = None) -> dict:
        """Process a base64 data URL: decode, validate, save to disk, run OCR.

        Returns a dict with keys: file_url, file_size, file_type, file_name, ocr_text (optional).
        Raises ValidationError on failure.
        """
        from django.core.files.base import ContentFile
        from django.core.files.storage import default_storage
        from common.upload_validators import validate_file_upload

        data: dict[str, object] = {}

        header, encoded = file_url.split(',', 1)
        mime_type = header.split(';')[0].split(':')[1] if ':' in header else file_type
        file_data = base64.b64decode(encoded)

        safe_name = file_name or f"upload-{document_identifier or 'pending'}"
        validate_file_upload(
            file_name=safe_name,
            mime_type=mime_type,
            file_bytes=file_data,
            field_name="file_url",
        )

        data["file_size"] = len(file_data)
        if not file_type and mime_type:
            data["file_type"] = mime_type
        if not file_name:
            data["file_name"] = safe_name

        doc_id = document_identifier or 'pending'
        media_root = settings.MEDIA_ROOT
        dms_dir = os.path.join(media_root, 'dms_versions', doc_id)
        os.makedirs(dms_dir, exist_ok=True)

        safe_filename = (data.get('file_name', safe_name) or safe_name).replace(' ', '_').replace('/', '_')
        file_path = os.path.join('dms_versions', doc_id, safe_filename)
        saved_path = default_storage.save(file_path, ContentFile(file_data, name=safe_filename))

        media_url = settings.MEDIA_URL or '/media/'
        if not media_url.startswith('/'):
            media_url = f'/{media_url}'
        data['file_url'] = f"{media_url.rstrip('/')}/{saved_path}"

        if mime_type in cls.SUPPORTED_OCR_TYPES:
            try:
                file_full_path = os.path.join(media_root, saved_path)
                ocr_text = OCRService.extract_text(file_full_path, mime_type)
                if ocr_text:
                    data['ocr_text'] = ocr_text
            except Exception as e:
                logger.warning("OCR extraction failed after upload: %s", e)

        return data


class DocumentAnalyticsService:
    """Service for document analytics and statistics."""
    
    @classmethod
    def get_document_stats(cls, document) -> dict:
        """Get statistics for a document."""
        versions = document.versions.all()
        permissions = document.permissions.all()
        comments = document.comments.all()
        access_logs = document.access_logs.all()
        
        # Calculate total views, downloads, and prints
        view_count = access_logs.filter(action='view').count()
        download_count = access_logs.filter(action='download').count()
        print_count = access_logs.filter(action='print').count()
        
        # Get unique viewers
        unique_viewers = access_logs.values('user').distinct().count()
        
        # Get word count from latest version
        latest_version = versions.order_by('-version_number').first()
        word_count = 0
        if latest_version:
            content = latest_version.content_text or latest_version.ocr_text or ""
            word_count = len(content.split())
        
        return {
            "version_count": versions.count(),
            "permission_count": permissions.count(),
            "comment_count": comments.count(),
            "unresolved_comment_count": comments.filter(resolved=False).count(),
            "view_count": view_count,
            "download_count": download_count,
            "print_count": print_count,
            "unique_viewers": unique_viewers,
            "word_count": word_count,
            "has_ocr_text": bool(latest_version and latest_version.ocr_text),
            "has_summary": bool(latest_version and latest_version.summary),
        }

