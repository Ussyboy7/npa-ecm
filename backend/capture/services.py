"""Content capture services for OCR, scanning, and batch processing."""

from __future__ import annotations

import logging
import os
import time
from typing import Any, Callable, Dict, List, Optional

import pytesseract
from django.conf import settings
from django.core.files.storage import default_storage
from pdf2image import convert_from_path
from PIL import Image
from docx import Document as DocxDocument

from dms.models import Document, DocumentVersion
from capture.models import CaptureJob

logger = logging.getLogger(__name__)


class OCRService:
    """Service for Optical Character Recognition processing."""

    @staticmethod
    def extract_text_from_image(image_path: str, language: str = "eng") -> Dict[str, Any]:
        """
        Extract text from an image using Tesseract OCR.

        Args:
            image_path: Path to the image file
            language: Language code for OCR (default: 'eng')

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            start_time = time.time()

            # Open and process image
            image = Image.open(image_path)

            # Perform OCR
            extracted_text = pytesseract.image_to_string(image, lang=language)

            # Get detailed data with confidence scores
            ocr_data = pytesseract.image_to_data(image, lang=language, output_type=pytesseract.Output.DICT)

            # Calculate average confidence
            confidences = [int(conf) for conf in ocr_data.get("conf", []) if conf != "-1"]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

            processing_time = time.time() - start_time

            return {
                "extracted_text": extracted_text.strip(),
                "confidence_score": avg_confidence / 100.0,  # Normalize to 0-1
                "language": language,
                "processing_time_seconds": processing_time,
                "word_count": len(extracted_text.split()),
            }
        except Exception as e:
            logger.error(f"OCR processing failed for {image_path}: {str(e)}")
            raise

    @staticmethod
    def extract_text_from_pdf(pdf_path: str, language: str = "eng", progress_callback: Optional[Callable[[int, int, int], None]] = None) -> Dict[str, Any]:
        """
        Extract text from a PDF file by converting pages to images and OCRing.

        Args:
            pdf_path: Path to the PDF file
            language: Language code for OCR
            progress_callback: Optional callback function(page_num, total_pages, progress_percent)

        Returns:
            Dictionary with extracted text, per-page results, and metadata
        """
        try:
            start_time = time.time()

            # Convert PDF pages to images
            # Note: Requires poppler-utils to be installed
            images = convert_from_path(pdf_path, dpi=300)
            total_pages = len(images)

            page_results = []
            full_text = []
            all_confidences = []

            for page_num, image in enumerate(images, 1):
                # Update progress if callback provided
                if progress_callback:
                    progress_percent = int((page_num / total_pages) * 100)
                    progress_callback(page_num, total_pages, progress_percent)

                # Save temporary image
                temp_image_path = f"/tmp/pdf_page_{page_num}.png"
                image.save(temp_image_path)

                try:
                    # Extract text from page
                    page_text = pytesseract.image_to_string(image, lang=language)
                    ocr_data = pytesseract.image_to_data(image, lang=language, output_type=pytesseract.Output.DICT)

                    # Calculate confidence
                    confidences = [int(conf) for conf in ocr_data.get("conf", []) if conf != "-1"]
                    page_confidence = sum(confidences) / len(confidences) if confidences else 0.0

                    page_results.append({
                        "page": page_num,
                        "text": page_text.strip(),
                        "confidence": page_confidence / 100.0,
                    })

                    full_text.append(f"--- Page {page_num} ---\n{page_text}")

                    if page_confidence > 0:
                        all_confidences.append(page_confidence)

                finally:
                    # Clean up temporary image
                    if os.path.exists(temp_image_path):
                        os.remove(temp_image_path)

            avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0.0
            processing_time = time.time() - start_time

            return {
                "extracted_text": "\n\n".join([r["text"] for r in page_results]),
                "full_text": "\n\n".join(full_text),
                "confidence_score": avg_confidence / 100.0,
                "language": language,
                "page_count": len(images),
                "page_results": page_results,
                "processing_time_seconds": processing_time,
                "word_count": sum(len(r["text"].split()) for r in page_results),
            }
        except Exception as e:
            logger.error(f"PDF OCR processing failed for {pdf_path}: {str(e)}")
            raise

    @staticmethod
    def extract_text_from_docx(docx_path: str) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            start_time = time.time()

            # Open the DOCX file
            doc = DocxDocument(docx_path)

            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_texts.append(" | ".join(row_text))

            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n--- Tables ---\n\n" + "\n".join(table_texts)

            processing_time = time.time() - start_time

            return {
                "extracted_text": full_text,
                "confidence_score": 1.0,  # DOCX text extraction is 100% accurate
                "language": "eng",  # Default, can be enhanced with language detection
                "processing_time_seconds": processing_time,
                "word_count": len(full_text.split()),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }
        except Exception as e:
            logger.error(f"DOCX text extraction failed for {docx_path}: {str(e)}")
            raise

    @staticmethod
    def process_document(document_id: str, language: str = "eng", force_reprocess: bool = False, progress_callback: Optional[Callable[[int, int, int], None]] = None) -> Dict[str, Any]:
        """
        Process a document for OCR extraction.

        Args:
            document_id: UUID of the document
            language: Language code for OCR
            force_reprocess: If True, reprocess even if OCR result exists

        Returns:
            OCR results dictionary
        """
        try:
            document = Document.objects.get(id=document_id)
            latest_version = document.versions.order_by("-version_number").first()

            if not latest_version:
                raise ValueError("Document has no versions")

            # Check for existing OCR result if not forcing reprocess
            if not force_reprocess:
                from capture.models import OCRResult
                existing_result = OCRResult.objects.filter(
                    document=document,
                    capture_job__status=CaptureJob.JobStatus.COMPLETED
                ).order_by('-created_at').first()
                
                if existing_result and latest_version.ocr_text:
                    logger.info(f"Using existing OCR result for document {document_id}")
                    return {
                        "extracted_text": existing_result.extracted_text,
                        "full_text": existing_result.full_text or existing_result.extracted_text,
                        "confidence_score": existing_result.confidence_score or 0.0,
                        "language": existing_result.language,
                        "page_count": existing_result.page_count,
                        "page_results": existing_result.page_results or [],
                        "processing_time_seconds": existing_result.processing_time_seconds or 0.0,
                        "word_count": len(existing_result.extracted_text.split()),
                        "from_cache": True,
                    }

            # Get file path - resolve from file_url
            file_url = latest_version.file_url
            if not file_url:
                raise ValueError("Document version has no file URL. Please upload a file first.")
            
            # Resolve file path similar to dms/views.py
            if file_url.startswith('/media/'):
                # Remove /media/ prefix and join with MEDIA_ROOT
                file_path = os.path.join(str(settings.MEDIA_ROOT), file_url.replace('/media/', ''))
            elif file_url.startswith(('http://', 'https://')):
                raise ValueError("Cannot process remote files for OCR. File must be stored locally.")
            else:
                # Assume it's a relative path from MEDIA_ROOT
                file_path = os.path.join(str(settings.MEDIA_ROOT), file_url.lstrip('/'))
            
            # Normalize the path (resolve any .. or .)
            file_path = os.path.normpath(file_path)
            
            # Verify file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found at path: {file_path} (resolved from file_url: {file_url})")

            # Determine file type and process accordingly
            # First try to get extension from file_name (more reliable)
            file_extension = None
            if latest_version.file_name:
                file_extension = os.path.splitext(latest_version.file_name)[1].lower()
            
            # If no extension from file_name, try file_path
            if not file_extension:
                file_extension = os.path.splitext(file_path)[1].lower()
            
            # If still no extension, try to infer from file_type (MIME type)
            if not file_extension and latest_version.file_type:
                mime_to_ext = {
                    'application/pdf': '.pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                    'application/msword': '.doc',
                    'image/png': '.png',
                    'image/jpeg': '.jpg',
                    'image/jpg': '.jpg',
                    'image/tiff': '.tiff',
                    'image/bmp': '.bmp',
                }
                file_extension = mime_to_ext.get(latest_version.file_type.lower(), '')
            
            # Validate file type before processing
            SUPPORTED_EXTENSIONS = [".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".docx"]
            SUPPORTED_MIME_TYPES = [
                'application/pdf',
                'image/png', 'image/jpeg', 'image/jpg',
                'image/tiff', 'image/bmp',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            ]
            
            if not file_extension:
                raise ValueError(
                    f"Cannot determine file type for '{latest_version.file_name}'. "
                    f"Supported formats: PDF, PNG, JPG, TIFF, BMP, DOCX. "
                    f"Please ensure the file has a valid extension."
                )
            
            if file_extension not in SUPPORTED_EXTENSIONS:
                raise ValueError(
                    f"Unsupported file type: {file_extension}. "
                    f"Supported formats: PDF, PNG, JPG, TIFF, BMP, DOCX. "
                    f"If this is a .doc file, please convert it to .docx format."
                )
            
            # Validate MIME type if available
            if latest_version.file_type and latest_version.file_type.lower() not in [m.lower() for m in SUPPORTED_MIME_TYPES]:
                logger.warning(
                    f"MIME type '{latest_version.file_type}' doesn't match extension '{file_extension}'. "
                    f"Proceeding with extension-based processing."
                )
            
            # Log for debugging
            logger.info(f"Processing document {document_id}: file_name={latest_version.file_name}, file_type={latest_version.file_type}, file_extension={file_extension}, file_path={file_path}")

            if file_extension == ".pdf":
                result = OCRService.extract_text_from_pdf(file_path, language, progress_callback)
            elif file_extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                result = OCRService.extract_text_from_image(file_path, language)
            elif file_extension == ".docx":
                result = OCRService.extract_text_from_docx(file_path)
            else:
                # This should not be reached due to validation above, but kept as safety
                raise ValueError(f"Unsupported file type for OCR: {file_extension}")

            return result
        except Document.DoesNotExist:
            raise ValueError(
                f"Document {document_id} not found. "
                f"Please ensure the document exists and you have permission to access it."
            )
        except FileNotFoundError as e:
            raise FileNotFoundError(
                f"File not found: {str(e)}. "
                f"The document file may have been moved or deleted. "
                f"Please re-upload the document."
            )
        except ValueError as e:
            # Re-raise ValueError as-is (already has good error messages)
            raise
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Document OCR processing failed for {document_id}: {error_msg}")
            
            # Provide more helpful error messages
            if "tesseract" in error_msg.lower() or "tesseract not found" in error_msg.lower():
                raise RuntimeError(
                    "OCR engine (Tesseract) is not installed or not found in PATH. "
                    "Please install Tesseract OCR: https://github.com/tesseract-ocr/tesseract"
                )
            elif "poppler" in error_msg.lower() or "pdftoppm" in error_msg.lower():
                raise RuntimeError(
                    "PDF processing requires Poppler utilities. "
                    "Please install poppler-utils: apt-get install poppler-utils (Linux) or brew install poppler (macOS)"
                )
            elif "permission denied" in error_msg.lower():
                raise PermissionError(
                    f"Permission denied accessing file. "
                    f"Please check file permissions and ensure the application has read access."
                )
            else:
                raise RuntimeError(
                    f"OCR processing failed: {error_msg}. "
                    f"If this persists, please contact support with the document details."
                )


class MetadataExtractionService:
    """Service for intelligent metadata extraction from documents."""

    @staticmethod
    def extract_metadata(text: str, document_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract metadata from document text using pattern matching.

        Args:
            text: Extracted text from OCR
            document_type: Type of document (optional)

        Returns:
            Dictionary with extracted metadata
        """
        import re

        metadata = {
            "dates": [],
            "reference_numbers": [],
            "names": [],
            "emails": [],
            "phone_numbers": [],
        }

        # Extract dates (various formats)
        date_patterns = [
            r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}",  # DD/MM/YYYY or DD-MM-YYYY
            r"\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4}",  # DD Month YYYY
        ]
        for pattern in date_patterns:
            dates = re.findall(pattern, text, re.IGNORECASE)
            metadata["dates"].extend(dates)

        # Extract reference numbers (patterns like REF: ABC123, Ref No: XYZ-456)
        ref_patterns = [
            r"(?:ref|reference|ref\s*no)[:\s]+([A-Z0-9\-/]+)",
            r"[A-Z]{2,10}[-/]\d{4,10}",
        ]
        for pattern in ref_patterns:
            refs = re.findall(pattern, text, re.IGNORECASE)
            metadata["reference_numbers"].extend(refs)

        # Extract emails
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        metadata["emails"] = re.findall(email_pattern, text)

        # Extract phone numbers (Nigerian format)
        phone_patterns = [
            r"\+?234[-\s]?\d{3}[-\s]?\d{3}[-\s]?\d{4}",  # +234 format
            r"0\d{3}[-\s]?\d{3}[-\s]?\d{4}",  # 0XXX format
        ]
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            metadata["phone_numbers"].extend(phones)

        # Remove duplicates
        for key in metadata:
            metadata[key] = list(set(metadata[key]))

        return metadata


class BatchProcessingService:
    """Service for batch document processing."""

    @staticmethod
    def process_batch(file_paths: List[str], process_ocr: bool = False, extract_metadata: bool = False) -> Dict[str, Any]:
        """
        Process multiple files in batch.

        Args:
            file_paths: List of file paths to process
            process_ocr: Whether to perform OCR
            extract_metadata: Whether to extract metadata

        Returns:
            Batch processing results
        """
        results = {
            "total": len(file_paths),
            "successful": 0,
            "failed": 0,
            "errors": [],
            "results": [],
        }

        for file_path in file_paths:
            try:
                file_result = {"file": os.path.basename(file_path), "status": "success"}

                if process_ocr:
                    # Determine file type and process
                    file_extension = os.path.splitext(file_path)[1].lower()
                    if file_extension == ".pdf":
                        ocr_result = OCRService.extract_text_from_pdf(file_path)
                    elif file_extension in [".png", ".jpg", ".jpeg"]:
                        ocr_result = OCRService.extract_text_from_image(file_path)
                    elif file_extension == ".docx":
                        ocr_result = OCRService.extract_text_from_docx(file_path)
                    else:
                        ocr_result = None

                    if ocr_result:
                        file_result["ocr"] = {
                            "text_length": len(ocr_result.get("extracted_text", "")),
                            "confidence": ocr_result.get("confidence_score", 0),
                        }

                        if extract_metadata and ocr_result.get("extracted_text"):
                            metadata = MetadataExtractionService.extract_metadata(
                                ocr_result["extracted_text"]
                            )
                            file_result["metadata"] = metadata

                results["results"].append(file_result)
                results["successful"] += 1

            except Exception as e:
                error_msg = f"Failed to process {file_path}: {str(e)}"
                results["errors"].append({"file": os.path.basename(file_path), "error": error_msg})
                results["failed"] += 1
                logger.error(error_msg)

        return results

